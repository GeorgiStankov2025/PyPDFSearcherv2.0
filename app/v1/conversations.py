import uuid
from datetime import datetime

from docx import Document
from fastapi import APIRouter, Depends, Response, HTTPException
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app import thread_variables
from app.db import Conversation, get_async_session, User, ReportRequest
from app.schemas import CreateReport
from app.files import download_file
from app.report_requests import create_report
from app.v1.users import verify_token, get_current_user

router=APIRouter(dependencies=[Depends(verify_token)])

@router.post("/conversation",tags=["conversations"])
async def create_conversation(session:AsyncSession=Depends(get_async_session),current_user: dict = Depends(get_current_user)):
    username_search = await session.execute(select(User).where(User.username == current_user["username"]))
    user = username_search.scalars().first()

    conversation=Conversation(

        user_id=user.id,
        topic ="",
        created_at =datetime.now(),
        modified_at =datetime.now(),

    )

    session.add(conversation)
    await session.commit()
    await session.refresh(conversation)
    return conversation

@router.get("/conversations",tags=["conversations"])
async def get_conversations(session:AsyncSession=Depends(get_async_session),current_user:dict=Depends(get_current_user)):

        user_search=await session.execute(select(User).where(User.username==current_user["username"]))
        user=user_search.scalars().first()
        conversations=await session.execute(select(Conversation).where(Conversation.user_id==user.id).order_by(Conversation.created_at))
        result= [row[0] for row in conversations.all()]
        return result

@router.get("/conversations/{conversation_id}",tags=["conversations"])
async def get_conversation(conversation_id:str,session:AsyncSession=Depends(get_async_session),current_user:dict=Depends(get_current_user)):

    try:
        user_search=await session.execute(select(User).where(User.username==current_user["username"]))
        user=user_search.scalars().first()
        conversation_uuid=uuid.UUID(conversation_id)
        conversation_search=await session.execute(select(Conversation).where(Conversation.id == conversation_uuid)
        .options(selectinload(Conversation.report_requests)))
        conversation=conversation_search.scalars().first()
        if conversation is None:
            raise HTTPException(status_code=404,detail="Conversation not found or deleted.")
        if user.id!=conversation.user_id:
            raise HTTPException(status_code=403,detail="You are not authorized to perform this action.")
        return conversation
    except HTTPException as e:
        raise HTTPException(status_code=e.status_code,detail=str(e))

@router.patch("/conversations/{conversation_id}",tags=["conversations"])
async def create_request_in_conversation(request:CreateReport,conversation_id:str,response:Response,session:AsyncSession=Depends(get_async_session),current_user:dict=Depends(get_current_user)):
    try:

        if request.message=="" or request.message is None:
            raise HTTPException(status_code=400,detail="Message cannot be empty")

        conversation_uuid=uuid.UUID(conversation_id)
        conversation_search = await session.execute(select(Conversation).where(Conversation.id == conversation_uuid))
        conversation = conversation_search.scalars().first()

        if conversation is None:
            raise HTTPException(status_code=404,detail="Conversation not found or deleted")
        username=current_user["username"]
        search = await session.execute(select(User).where(User.username ==username))
        user = search.scalars().first()

        if user is None:
            raise HTTPException(status_code=404,detail="User not found or deleted")
        if user.id!=conversation.user_id:
            raise HTTPException(status_code=403,detail="You are not authorized to perform this action!")

        if conversation.topic=="":
            conversation.topic=request.message

        thread_variables.topic=conversation.topic
        document_content=await create_report(request,response,username,conversation_id)

        report_request = ReportRequest(

            conversation_id=conversation.id,
            input_message=request.message,
            response=document_content,
            created_at=datetime.now(),

        )
        conversation.modified_at=datetime.now()
        if 'I cannot fulfill this request' in document_content:
            session.add(report_request)
            await session.commit()
            await session.refresh(report_request)
            return "I cannot fulfill this request because the required information is not present in the database. Try to be more specific or choose a different topic."
        else:
            report_request.is_successful=True
            document_content = document_content.replace("##", "")
            document_title = document_content.splitlines()[0]
            document = Document()
            document.add_heading(document_title, level=1)
            document.add_heading(f"Author:{current_user['username']}", level=2)
            document.add_paragraph(document_content)
            document.save(rf"D:\ПУ\II курс\Python\PyPDFSearcher\generated_reports\{document_title}.docx")
            session.add(report_request)
            await session.commit()
            await session.refresh(report_request)
            return await download_file(document_title + ".docx")
    except HTTPException as e:
        raise HTTPException(status_code=e.status_code,detail=str(e))

@router.delete("/conversations/{conversation_id}",tags=["conversations"])
async def delete_conversation(conversation_id:str,session:AsyncSession=Depends(get_async_session),current_user:dict=Depends(get_current_user)):

    try:
        user_search = await session.execute(select(User).where(User.username == current_user["username"]))
        user = user_search.scalars().first()
        conversation_uuid = uuid.UUID(conversation_id)
        conversation_search = await session.execute(select(Conversation).where(Conversation.id ==conversation_uuid))
        conversation = conversation_search.scalars().first()
        if conversation is None:
            raise HTTPException(status_code=404,detail="Conversation not found or deleted")
        if user.id!=conversation.user_id:
            raise HTTPException(status_code=403,detail="You are not authorized to perform this action.")
        await session.delete(conversation)
        await session.commit()
        return "Conversation deleted!"
    except HTTPException as e:
        raise HTTPException(status_code=e.status_code,detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500,detail=str(e))